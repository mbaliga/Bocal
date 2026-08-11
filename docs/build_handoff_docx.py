#!/usr/bin/env python3
"""Build the Bocal handoff DOCX from reviewed Markdown sources."""

from __future__ import annotations

import csv
import re
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = Path(__file__).resolve().parent / "Bocal_Product_Handoff.docx"
SOURCES = [
    (Path(__file__).resolve().parent / "Bocal_Product_Handoff.md", False),
    (ROOT / "research" / "Personas_and_50_Workflows.md", True),
    (ROOT / "research" / "Research_Sources_and_Method.md", True),
]

INK = "182033"
BLUE = "274A78"
VIOLET = "6959D9"
CYAN = "0A8F83"
PALE = "E8EEF5"
PALE_VIOLET = "F1EFFF"
RULE = "C7D2E0"
MUTED = "5D6777"
WHITE = "FFFFFF"


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=110, bottom=80, end=110) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, **edges) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge, options in edges.items():
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        for key, value in options.items():
            element.set(qn(f"w:{key}"), str(value))


def no_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant = OxmlElement("w:cantSplit")
    tr_pr.append(cant)


def repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_hyperlink(paragraph, label: str, url: str) -> None:
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), VIOLET)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.extend([color, underline])
    text = OxmlElement("w:t")
    text.text = label
    run.extend([properties, text])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


INLINE_RE = re.compile(r"(\[.+?\]\(https?://[^)]+\)|\*\*.+?\*\*|`[^`]+`)")


def add_inline(paragraph, text: str) -> None:
    position = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > position:
            paragraph.add_run(text[position : match.start()])
        token = match.group(0)
        if token.startswith("["):
            parsed = re.match(r"\[(.+?)\]\((https?://[^)]+)\)", token)
            if parsed:
                add_hyperlink(paragraph, parsed.group(1), parsed.group(2))
        elif token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.2)
            run.font.color.rgb = RGBColor.from_string(BLUE)
        position = match.end()
    if position < len(text):
        paragraph.add_run(text[position:])


def set_run_defaults(run, size=10.3, color=INK) -> None:
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)


def add_page_number(paragraph) -> None:
    paragraph.add_run("PAGE ")
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, end])


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.83)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)
    section.header_distance = Inches(0.33)
    section.footer_distance = Inches(0.33)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.3)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(5.5)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, color, before, after in (
        ("Title", 34, INK, 0, 10),
        ("Heading 1", 19, BLUE, 17, 8),
        ("Heading 2", 14, BLUE, 14, 6),
        ("Heading 3", 11.8, INK, 10, 4),
    ):
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(10.1)
        style.paragraph_format.left_indent = Inches(0.25)
        style.paragraph_format.first_line_indent = Inches(-0.16)
        style.paragraph_format.space_after = Pt(3)

    for sec in document.sections:
        header = sec.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        left = hp.add_run("BOCAL  /  PRODUCT HANDOFF")
        left.bold = True
        left.font.color.rgb = RGBColor.from_string(BLUE)
        left.font.size = Pt(8.5)
        hp.add_run("\t10 AUGUST 2026").font.size = Pt(8.5)
        hp.paragraph_format.tab_stops.add_tab_stop(Inches(6.4))
        footer = sec.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        fp.add_run("LOCAL-FIRST MUSIC LEARNING  ·  ")
        add_page_number(fp)
        for run in fp.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor.from_string(MUTED)


def add_cover(document: Document) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(65)
    run = p.add_run("BOCAL")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor.from_string(VIOLET)
    run.font.letter_spacing = Pt(1.4)

    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(12)
    title.paragraph_format.keep_with_next = True
    r = title.add_run("Music practice,\nmade legible.")
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(35)
    r.font.color.rgb = RGBColor.from_string(INK)

    line_table = document.add_table(rows=1, cols=2)
    line_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    line_table.autofit = False
    line_table.columns[0].width = Inches(1.4)
    line_table.columns[1].width = Inches(5.3)
    shade(line_table.cell(0, 0), VIOLET)
    shade(line_table.cell(0, 1), CYAN)
    for cell in line_table.rows[0].cells:
        cell.height = Inches(0.08)
        set_cell_margins(cell, 0, 0, 0, 0)

    sub = document.add_paragraph()
    sub.paragraph_format.space_before = Pt(22)
    sub.paragraph_format.space_after = Pt(28)
    sr = sub.add_run("Product strategy · competitor parity · 50 persona workflows\nNative Android source · static web build · 35 interactive woodwind models")
    sr.font.size = Pt(16)
    sr.font.color.rgb = RGBColor.from_string(BLUE)
    sr.font.bold = True

    card = document.add_table(rows=1, cols=1)
    card.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = card.cell(0, 0)
    shade(cell, PALE_VIOLET)
    set_cell_margins(cell, 180, 180, 180, 180)
    cp = cell.paragraphs[0]
    cr = cp.add_run("HANDOFF STATUS")
    cr.bold = True
    cr.font.size = Pt(9)
    cr.font.color.rgb = RGBColor.from_string(VIOLET)
    cp = cell.add_paragraph()
    add_inline(cp, "Working static reference and structurally validated assets. Native source supplied. No fake APK; physical Android compilation and specialist fingering sign-off remain explicit gates.")

    date = document.add_paragraph()
    date.paragraph_format.space_before = Pt(54)
    date.add_run("REFERENCE HANDOFF  /  10 AUGUST 2026").font.color.rgb = RGBColor.from_string(MUTED)
    date.runs[0].font.size = Pt(9)
    date.runs[0].bold = True
    document.add_page_break()


def add_contents(document: Document) -> None:
    document.add_heading("Contents", level=1)
    entries = [
        "Executive decision",
        "1. Product thesis",
        "2. Research and incumbent definition",
        "3. Personas and user journeys",
        "4. Information architecture and UX",
        "5. Delivered application capability",
        "6. 3D asset system",
        "7. Technical architecture",
        "8. Parity status and roadmap",
        "9. Quality and release plan",
        "10. Delta from the supplied baseline",
        "11. Delivery map",
        "12. Immediate next decisions",
        "Appendix A. Ten personas and 50 workflows",
        "Appendix B. Research sources and method",
    ]
    for index, entry in enumerate(entries, 1):
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        num = p.add_run(f"{index:02d}  ")
        num.bold = True
        num.font.color.rgb = RGBColor.from_string(VIOLET)
        add_inline(p, entry)
    note = document.add_paragraph()
    note.paragraph_format.space_before = Pt(18)
    add_inline(note, "Companion files: 84-row TE parity CSV, detailed baseline delta, source code, ready static build, 35-model pack and native Android Studio project.")
    document.add_page_break()


def add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    columns = max(len(row) for row in rows)
    table = document.add_table(rows=0, cols=columns)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for row_index, values in enumerate(rows):
        row = table.add_row()
        no_row_split(row)
        if row_index == 0:
            repeat_header(row)
        for column_index in range(columns):
            cell = row.cells[column_index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            set_cell_border(
                cell,
                top={"val": "single", "sz": 4, "color": RULE},
                bottom={"val": "single", "sz": 4, "color": RULE},
                start={"val": "single", "sz": 4, "color": RULE},
                end={"val": "single", "sz": 4, "color": RULE},
            )
            if row_index == 0:
                shade(cell, PALE)
            elif row_index % 2 == 0:
                shade(cell, "F7F9FC")
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            add_inline(paragraph, values[column_index] if column_index < len(values) else "")
            for run in paragraph.runs:
                set_run_defaults(run, 8.2 if columns >= 4 else 8.8, INK)
                if row_index == 0:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(BLUE)
    document.add_paragraph().paragraph_format.space_after = Pt(1)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        values = [value.strip() for value in lines[index].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", value) for value in values):
            rows.append(values)
        index += 1
    return rows, index


def add_markdown(document: Document, path: Path, appendix: bool = False) -> None:
    lines = path.read_text(encoding="utf-8").splitlines()
    if appendix:
        document.add_page_break()
    index = 0
    in_code = False
    code_lines: list[str] = []
    skip_first_h1 = not appendix
    while index < len(lines):
        raw = lines[index]
        stripped = raw.strip()
        if stripped.startswith("```"):
            if in_code:
                p = document.add_paragraph()
                p.style = document.styles["Normal"]
                p.paragraph_format.left_indent = Inches(0.18)
                p.paragraph_format.right_indent = Inches(0.18)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(7)
                p.paragraph_format.keep_together = True
                text = "\n".join(code_lines)
                if raw.strip() == "```" and code_lines and any("-->" in line for line in code_lines):
                    text = "Flow: " + "  →  ".join(re.findall(r'\["(.+?)"\]', " ".join(code_lines)))
                run = p.add_run(text)
                run.font.name = "Consolas"
                run.font.size = Pt(8.2)
                run.font.color.rgb = RGBColor.from_string(BLUE)
                pPr = p._p.get_or_add_pPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), "F3F5F8")
                pPr.append(shd)
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(raw)
            index += 1
            continue
        if not stripped:
            index += 1
            continue
        if stripped.startswith("|"):
            rows, index = parse_table(lines, index)
            add_table(document, rows)
            continue
        heading = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if heading:
            level = len(heading.group(1))
            text = heading.group(2)
            if level == 1 and skip_first_h1:
                skip_first_h1 = False
                index += 1
                continue
            if appendix and level == 1:
                prefix = "Appendix A" if "personas" in text.lower() else "Appendix B"
                text = f"{prefix}. {text}"
            paragraph = document.add_heading(level=min(3, level), text="")
            add_inline(paragraph, text)
            index += 1
            continue
        if stripped.startswith("> "):
            card = document.add_table(rows=1, cols=1)
            card.alignment = WD_TABLE_ALIGNMENT.LEFT
            cell = card.cell(0, 0)
            shade(cell, PALE_VIOLET)
            set_cell_margins(cell, 130, 150, 130, 150)
            p = cell.paragraphs[0]
            add_inline(p, stripped[2:])
            index += 1
            continue
        bullet = re.match(r"^[-*]\s+(.+)$", stripped)
        numbered = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        if bullet or numbered:
            if bullet:
                p = document.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.28)
                p.paragraph_format.first_line_indent = Inches(-0.18)
                p.paragraph_format.space_after = Pt(3)
                marker = p.add_run("•  ")
                marker.bold = True
                marker.font.color.rgb = RGBColor.from_string(VIOLET)
                add_inline(p, bullet.group(1))
            else:
                p = document.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.28)
                p.paragraph_format.first_line_indent = Inches(-0.22)
                p.paragraph_format.space_after = Pt(3)
                marker = p.add_run(f"{numbered.group(1)}.  ")
                marker.bold = True
                marker.font.color.rgb = RGBColor.from_string(VIOLET)
                add_inline(p, numbered.group(2))
            index += 1
            continue
        p = document.add_paragraph()
        add_inline(p, stripped)
        index += 1


def add_parity_summary(document: Document) -> None:
    csv_path = ROOT / "research" / "TE_Parity_Matrix.csv"
    with csv_path.open(encoding="utf-8") as handle:
        rows = list(csv.DictReader(handle))
    document.add_page_break()
    document.add_heading("Companion parity matrix at a glance", level=1)
    p = document.add_paragraph()
    add_inline(p, f"The companion CSV contains {len(rows)} capability rows. It is the testable scope ledger; this summary shows row count by domain.")
    counts: dict[str, int] = {}
    for row in rows:
        counts[row["Domain"]] = counts.get(row["Domain"], 0) + 1
    add_table(document, [["Domain", "Capability rows"]] + [[key, str(value)] for key, value in counts.items()])


def main() -> None:
    document = Document()
    configure_document(document)
    add_cover(document)
    add_contents(document)
    for path, appendix in SOURCES:
        add_markdown(document, path, appendix=appendix)
    add_parity_summary(document)

    core = document.core_properties
    core.title = "Bocal product, UX, technical and asset handoff"
    core.subject = "Music learning product handoff and TonalEnergy parity program"
    core.author = "Bocal project"
    core.keywords = "Bocal, music learning, woodwind, saxophone, Android, tuner, metronome, glTF"
    core.comments = "Reference handoff dated 10 August 2026"
    document.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
